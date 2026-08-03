"""
研究計画書（DOCX, PDF, TXT）のテキスト抽出および倫理指針Wordファイルのパースモジュール
"""
import io
import os
import docx
import pypdf
from typing import Dict, List, Any

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    アップロードされたファイルまたはローカルファイルのバイトデータからテキストを抽出する。
    """
    ext = filename.lower().split('.')[-1]
    
    if ext == 'docx':
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text)
        
        # テーブル内のテキストも抽出
        for table in doc.tables:
            for row in table.rows:
                row_str = " | ".join([c.text.strip().replace('\n', ' ') for c in row.cells if c.text.strip()])
                if row_str:
                    full_text.append(row_str)
                    
        return "\n".join(full_text)
        
    elif ext == 'pdf':
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        full_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                full_text.append(t)
        return "\n".join(full_text)
        
    elif ext in ['txt', 'md']:
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            return file_bytes.decode('shift_jis', errors='ignore')
            
    else:
        raise ValueError(f"未対応のファイルフォーマットです: .{ext}")


def load_ethics_guideline_from_docx(docx_path: str = '人を対象とする生命科学・医学系研究に関する倫理指針.docx') -> list:
    """
    倫理指針Wordファイルを読み込み、章・節・項の階層構造を解析して返す。
    見出しスタイル (Heading 1 / Heading 2 / Heading 3) を基準に分割する。
    """
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        abs_path = docx_path if os.path.isabs(docx_path) else os.path.join(base_dir, docx_path)

        doc = docx.Document(abs_path)
        sections = []
        current = None
        body_lines = []

        HEADING_STYLES = {'Heading 1', 'Heading 2', 'Heading 3', 'Title'}

        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ''
            text = p.text.strip()
            if not text:
                continue

            if style_name in HEADING_STYLES:
                # 直前のセクションを保存
                if current is not None:
                    current['content'] = '\n'.join(body_lines)
                    sections.append(current)
                    body_lines = []

                level = {'Title': 0, 'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3}.get(style_name, 2)
                current = {
                    'title': text,
                    'level': level,
                    'content': '',
                    'style': style_name
                }
            else:
                if current is not None:
                    body_lines.append(text)

        # 最後のセクションを保存
        if current is not None:
            current['content'] = '\n'.join(body_lines)
            sections.append(current)

        return sections

    except Exception as e:
        return [{'title': 'エラー', 'level': 1,
                 'content': f'倫理指針Wordファイルの読み込みに失敗しました: {e}', 'style': 'Heading 1'}]


